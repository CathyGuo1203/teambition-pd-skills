#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 验证报告转换为 Word 格式
"""

import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_word_report(md_file_path, output_path=None):
    """读取 Markdown 报告并生成 Word 文档"""
    
    # 读取 Markdown 文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    # 添加标题
    title = doc.add_heading('项目计划验证报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 解析 Markdown 内容（简化版）
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line or line.startswith('---'):
            continue
        
        # 处理不同层级的标题
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('|') and line.endswith('|'):
            # 表格行（简化处理，作为文本添加）
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            paragraph = doc.add_paragraph()
            runner = paragraph.add_run('  '.join(cells))
            runner.font.size = Pt(9)
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            paragraph = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # 普通段落
            if line:
                doc.add_paragraph(line)
    
    # 保存 Word 文档
    if output_path is None:
        base_name = os.path.splitext(md_file_path)[0]
        output_path = base_name + '.docx'
    
    doc.save(output_path)
    print(f"✅ Word 报告已保存：{output_path}")
    
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python generate_word_report.py <Markdown 报告文件> [输出路径]")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else None
    
    create_word_report(md_file, output)
